
"""
Public portfolio version of invoice views.

NOTE:
This module is intentionally simplified.

The production system includes:
- advanced filtering and aggregation
- transactional operations
- audit logging
- complex business rules (inventory lifecycle, reporting)

These have been removed to protect proprietary logic
while preserving overall architecture and structure.
"""

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Q, Sum, F, Value, DecimalField
from django.db.models.functions import Coalesce
from django.core.paginator import Paginator
from django.http import HttpResponse
from decimal import Decimal
import io
 
from .models import Invoice, InvoiceItem, Payment
from .forms import InvoiceGenerateForm, InvoiceFilterForm, InvoiceItemForm, PaymentForm
from apps.operations.models import SellOrder
from apps.accounts.permissions import write_required, supervisor_required
from apps.audit.helpers import log_action
from apps.audit.models import AuditLog
from django.contrib.staticfiles import finders

 
 
@login_required
def invoice_list(request):
    form = InvoiceFilterForm(request.GET)
 
    # Annotate with paid_amount so we can filter by payment status in the DB
    qs = Invoice.objects.select_related('retailer').annotate(
        paid_ann=Coalesce(Sum('payments__amount'), Value(Decimal('0.00')),
                          output_field=DecimalField())
    ).order_by('-date', '-pk')
 
    if form.is_valid():
        #filtering logic here
 
    page_obj = Paginator(qs, 30).get_page(request.GET.get('page'))
    return render(request, 'invoices/invoice_list.html', {
        'page_obj': page_obj, 'form': form, 'total': qs.count(),
    })
 
 
@login_required
def invoice_detail(request, pk):
    invoice = get_object_or_404(
        Invoice.objects.prefetch_related('items', 'payments'), pk=pk
    )
    edit_requested = request.GET.get('edit') == '1'
    payment_form   = PaymentForm()
    return render(request, 'invoices/invoice_detail.html', {
        'invoice':        invoice,
        'edit_requested': edit_requested,
        'payment_form':   payment_form,
    })
 
 
@login_required
@write_required
def payment_add(request, invoice_pk):
    invoice = get_object_or_404(Invoice, pk=invoice_pk)
    if request.method == 'POST':
        form = PaymentForm(request.POST)
        if form.is_valid():
            payment = form.save(commit=False)
            payment.invoice = invoice
            payment.save()
            log_action(request, AuditLog.ACTION_CREATE, payment)
            messages.success(request, f'Payment of ${payment.amount} recorded.')
        else:
            messages.error(request, 'Invalid payment data. Please check the form.')
    return redirect('invoices:invoice_detail', pk=invoice_pk)
 
 
@login_required
@supervisor_required
def payment_delete(request, payment_pk):
    payment    = get_object_or_404(Payment, pk=payment_pk)
    invoice_pk = payment.invoice_id
    if request.method == 'POST':
        log_action(request, AuditLog.ACTION_DELETE, payment)
        payment.delete()
        messages.success(request, 'Payment removed.')
    return redirect('invoices:invoice_detail', pk=invoice_pk)
 
 
@login_required
@write_required
def generate_from_order(request, order_pk):
    order = get_object_or_404(SellOrder, pk=order_pk)
 
    if hasattr(order, 'invoice'):
        messages.info(request, f'Invoice already exists: {order.invoice.invoice_number}')
        return redirect('invoices:invoice_detail', pk=order.invoice.pk)
 
    form = InvoiceGenerateForm(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        discount = form.cleaned_data.get('discount') or Decimal('0.00')
        notice   = form.cleaned_data.get('extra_notice') or ''
        try:
            invoice = Invoice.generate_from_sell_order(order, discount=discount, extra_notice=notice)
            log_action(request, AuditLog.ACTION_CREATE, invoice)
            messages.success(request, f'Invoice {invoice.invoice_number} generated.')
            return redirect('invoices:invoice_detail', pk=invoice.pk)
        except ValueError as e:
            messages.error(request, str(e))
 
    return render(request, 'invoices/generate_form.html', {
        'order': order, 'form': form,
    })
 
 
@login_required
@write_required
def invoice_update(request, pk):
    invoice = get_object_or_404(Invoice, pk=pk)
    if request.method == 'POST':
        #your code
 
        # Update items
        
        invoice.update_total()
        log_action(request, AuditLog.ACTION_UPDATE, invoice)
        messages.success(request, f'Invoice {invoice.invoice_number} updated.')
        return redirect('invoices:invoice_detail', pk=pk)
    return redirect('invoices:invoice_detail', pk=pk)
 
 
@login_required
@supervisor_required
def invoice_delete(request, pk):
    invoice = get_object_or_404(Invoice, pk=pk)
    if request.method == 'POST':
        log_action(request, AuditLog.ACTION_DELETE, invoice)
        invoice.delete()
        messages.success(request, 'Invoice deleted.')
        return redirect('invoices:invoice_list')
    return render(request, 'invoices/confirm_delete.html', {
        'object': invoice, 'object_type': 'Invoice',
        'back_url': 'invoices:invoice_detail', 'back_pk': pk,
    })
 
@login_required
def invoice_export_docx(request, pk):
    """Generate and download the invoice as a .docx file."""
    invoice = get_object_or_404(Invoice.objects.prefetch_related('items'), pk=pk)
    log_action(request, AuditLog.ACTION_EXPORT, invoice)

    buffer = _build_docx(invoice)
    filename = f"{invoice.invoice_number}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _build_docx(invoice):
    """Build invoice .docx — aggregated, no bottle numbers."""
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def set_col_width(cell, width_cm):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcW')):
            tcPr.remove(old)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Logo (left aligned)

    # Header: company name

    # Company info + Invoice info side by side

    # Bill To
    doc.add_paragraph(f"Bill To: {invoice.org_name_snapshot}")
    doc.add_paragraph(f"Address: {invoice.org_address_snapshot}")
    #doc.add_paragraph(f"License: {invoice.license_snapshot}")
    doc.add_paragraph()

    # Items table — 4 columns, NO bottle numbers
    col_widths = [9, 2, 3, 3.5]
    headers    = ["Product", "Qty", "Unit Price", "Total"]

    items_table = doc.add_table(rows=1, cols=4)
    items_table.autofit = False

    # Fixed layout

    # Header row
    hrow = items_table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_col_width(cell, col_widths[i])

    # Data rows
    doc.add_paragraph()

    # Totals

    doc.add_paragraph("\n")
    sig = doc.add_paragraph("Signature: __________________________")
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
